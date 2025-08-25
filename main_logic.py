from docx import Document
from pypdf import PdfReader
from openai import OpenAI
from dotenv import load_dotenv
import os
import json

####### CARREGAR AS VARIÁVEIS DO ARQUIVO ENV ###########################
load_dotenv()

####### FUNÇÃO PRA EXTRAIR E PADRONIZAR A RESPOSTA DA IA ################
def limpar_e_extrair_json(resposta_ia):
    """
    Extrai uma string JSON de uma resposta de texto maior, ignorando texto adicional.
    """
    try:
        inicio_json = resposta_ia.find('{')
        fim_json = resposta_ia.rfind('}')
        
        if inicio_json != -1 and fim_json != -1 and fim_json > inicio_json:
            json_str = resposta_ia[inicio_json : fim_json + 1]
            return json_str
        else:
            return None
    except Exception as e:
        print(f"Erro ao tentar limpar a resposta da IA: {e}")
        return None

####### FUNÇÃO PARA PREENCHER O DOCX COM A RESPOSTA PARDRONIZADA ################

def preencher_documento(dados, caminho_template, caminho_saida):
    """
    Preenche um documento .docx a partir de um template.
    Esta versão procura os placeholders em parágrafos normais E dentro de tabelas.
    """
    try:
        modelo = Document(caminho_template)
        print(f"\n--- Iniciando preenchimento do template: '{caminho_template}' ---")

        chaves_utilizadas = {chave: False for chave in dados}

        for paragrafo in modelo.paragraphs:
            for chave, valor in dados.items():
                substituto = f'[{chave}]'
                if substituto in paragrafo.text:
                    valor_str = str(valor)
                    print(f"-> Encontrado '{substituto}' no corpo do texto. Substituindo...")
                    paragrafo.text = paragrafo.text.replace(substituto, valor_str)
                    chaves_utilizadas[chave] = True

        for tabela in modelo.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        for chave, valor in dados.items():
                            substituto = f'[{chave}]'
                            if substituto in paragrafo.text:
                                valor_str = str(valor)
                                print(f"-> Encontrado '{substituto}' em uma tabela. Substituindo...")
                                paragrafo.text = paragrafo.text.replace(substituto, valor_str)
                                chaves_utilizadas[chave] = True
        
        modelo.save(caminho_saida)
        print(f"\n--- SUCESSO ---")
        print(f"Documento final salvo em: '{caminho_saida}'")

        chaves_nao_encontradas = [chave for chave, utilizada in chaves_utilizadas.items() if not utilizada]
        if chaves_nao_encontradas:
            print("\nAVISO: As seguintes chaves do JSON não foram encontradas como placeholders no documento:")
            for chave in chaves_nao_encontradas:
                print(f"- {chave}")

    except FileNotFoundError:
        print(f"\n--- ERRO CRÍTICO ---")
        print(f"O arquivo de template '{caminho_template}' não foi encontrado. Verifique o nome e o caminho.")
    except Exception as e:
        print(f"\n--- ERRO CRÍTICO ---")
        print(f"Ocorreu um erro ao preencher ou salvar o documento: {e}")


####### EXTRAIR O TEXTO DO PDF QUE SERÁ ENVIADO PARA A IA ################

def extrair_texto(caminho_pdf):
    try:
        ler = PdfReader(caminho_pdf)
        edital_lido = ''
        for pagina in ler.pages:
            texto_pagina = pagina.extract_text()
            if texto_pagina:
                edital_lido += texto_pagina + '\n'
        return edital_lido
    except FileNotFoundError:
        print(f"Erro: O arquivo PDF não foi encontrado em '{caminho_pdf}'")
        return None
    except Exception as e:
        print(f"Erro: Ocorreu um erro ao extrair o texto do PDF: {e}")
        return None

####### FUNÇÃO ENVIAR PARA A IA ################

def analisar_edital(caminho_pdf):
    print("Iniciando a análise do edital. Isso pode levar alguns momentos...")
    texto_edital = extrair_texto(caminho_pdf)

    if not texto_edital:
        print("Não foi possível extrair texto do PDF. Abortando a análise.")
        return None

    try:
        with open('contexto.txt', 'r', encoding="utf-8") as f:
            contexto_completo = f.read()
    except FileNotFoundError:
        print("Erro: Arquivo 'contexto.txt' não encontrado. Verifique se ele está na mesma pasta do script.")
        return None

    client = OpenAI(
        api_key=os.getenv("DEEPSEEK_API_KEY"), 
        base_url="https://api.deepseek.com"
    )

    print("Enviando requisição para a IA...")
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": contexto_completo},
                {"role": "user", "content": f"Analise o edital a seguir e retorne os dados em formato JSON. \n--- INÍCIO DO EDITAL ---\n{texto_edital}\n--- FIM DO EDITAL ---"}
            ],
            max_tokens=4096, 
            temperature=0.3,
            stream=False
        )
        print("Requisição concluída com sucesso!")
        return response.choices[0].message.content
    except Exception as e:
        print(f"Erro ao fazer a requisição para a API: {e}")
        return None

####### FUNÇÃO START DE TUDO ################

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    caminho_edital_pdf = os.path.join(script_dir, 'EDITAL.pdf')
    caminho_modelo_docx = os.path.join(script_dir, 'RESUMO_SMP..docx') # Corrigi o nome, verifique se está certo
    caminho_arquivo_final = os.path.join(script_dir, 'Edital Analisado - Final.docx')

    resposta_bruta_ia = analisar_edital(caminho_edital_pdf)

    if resposta_bruta_ia:
        print("\n--- Resposta Bruta da IA Recebida ---")
        json_limpo = limpar_e_extrair_json(resposta_bruta_ia)
        
        if json_limpo:
            try:
                dados_edital = json.loads(json_limpo)
                print("\n--- JSON recebido e processado ---")
                print(json.dumps(dados_edital, indent=2, ensure_ascii=False)) # Imprime o JSON de forma legível

                preencher_documento(dados_edital, caminho_modelo_docx, caminho_arquivo_final)

            except json.JSONDecodeError:
                print("\n--- ERRO DE DECODIFICAÇÃO ---")
                print("A resposta da IA, mesmo após a limpeza, não é um JSON válido. Resposta limpa:")
                print(json_limpo)
        else:
            print("\n--- ERRO ---")
            print("Não foi possível extrair um JSON da resposta da IA. Resposta crua:")
            print(resposta_bruta_ia)

####### BLOCO DE EXEC ################
if __name__ == "__main__":
    main()