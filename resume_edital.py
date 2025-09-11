from docx import Document
from docx.shared import RGBColor
from pypdf import PdfReader
from openai import OpenAI
from dotenv import load_dotenv
import os
import json

load_dotenv()

def limpar_e_extrair_json(resposta_ia):
    try:
        inicio_json = resposta_ia.find('{')
        fim_json = resposta_ia.rfind('}')
        
        if inicio_json != -1 and fim_json != -1 and fim_json > inicio_json:
            json_str = resposta_ia[inicio_json : fim_json + 1]
            return json_str
        else:
            return None
    except Exception as e:
        print(f"Erro ao tentar limpar a resposta da IA: {e}")
        return None

def puxar_paragrafo(documento):
    for paragrafo in documento.paragraphs:
        yield paragrafo
    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    yield paragrafo

def substituir_formatar(documento, dados):
    print("#### Iniciando Substituição ####")
    for paragrafo in puxar_paragrafo(documento):
        for chave, valor in dados.items():
            substituto = f'[{chave}]'
            texto_formatado = ''.join(run.text for run in paragrafo.runs)

            if substituto in texto_formatado:
                runs = paragrafo.runs
                idx_inicio, idx_fim = -1, -1
                texto_temp = ''

                for i, run in enumerate(runs): 
                    texto_temp += run.text
                    if substituto in texto_temp:
                        if idx_inicio == -1:
                            idx_inicio = i
                        idx_fim = i
                        break

                if idx_inicio != -1:
                    run_alvo = runs[idx_inicio]
                    run_alvo.text = str(valor)

                    run_alvo.bold = True
                    run_alvo.font.color.rgb = RGBColor(0, 0, 139)

                    for i in range(idx_inicio + 1, idx_fim + 1):
                        runs[i].text = ''

                    print(f"-> Substituído '{substituto}' com formatação.")


def extrair_texto(caminho_pdf):
    try:
        ler = PdfReader(caminho_pdf)
        edital_lido = ''
        for pagina in ler.pages:
            texto_pagina = pagina.extract_text()
            if texto_pagina:
                edital_lido += texto_pagina + '\n'
        return edital_lido
    except FileNotFoundError:
        print(f"Erro: O arquivo PDF não foi encontrado em '{caminho_pdf}'")
        return None
    except Exception as e:
        print(f"Erro: Ocorreu um erro ao extrair o texto do PDF: {e}")
        return None

def analisar_edital(caminho_pdf):
    print("Iniciando a análise do edital. Isso pode levar alguns momentos...")
    texto_edital = extrair_texto(caminho_pdf)

    if not texto_edital:
        print("Não foi possível extrair texto do PDF. Abortando a análise.")
        return None

    try:
        with open('contexto.txt', 'r', encoding="utf-8") as f:
            contexto_completo = f.read()
    except FileNotFoundError:
        print("Erro: Arquivo 'contexto.txt' não encontrado. Verifique se ele está na mesma pasta do script.")
        return None

    client = OpenAI(
        api_key=os.getenv("DEEPSEEK_API_KEY"), 
        base_url="https://api.deepseek.com"
    )

    print("Enviando requisição para a IA...")
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": contexto_completo},
                {"role": "user", "content": f"Analise o edital a seguir e retorne os dados em formato JSON. \n--- INÍCIO DO EDITAL ---\n{texto_edital}\n--- FIM DO EDITAL ---"}
            ],
            max_tokens=4096, 
            temperature=0.3,
            stream=False
        )
        print("Requisição concluída com sucesso!")
        return response.choices[0].message.content
    except Exception as e:
        print(f"Erro ao fazer a requisição para a API: {e}")
        return None

def main():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    caminho_edital_pdf = os.path.join(script_dir, 'EDITAL.pdf')
    caminho_modelo_docx = os.path.join(script_dir, 'RESUMO_SMP..docx')
    caminho_arquivo_final = os.path.join(script_dir, 'Edital Analisado - Final.docx')

    resposta_bruta_ia = analisar_edital(caminho_edital_pdf)

    if resposta_bruta_ia:
        print("\n--- Resposta Bruta da IA Recebida ---")
        json_limpo = limpar_e_extrair_json(resposta_bruta_ia)
        
        if json_limpo:
            try:
                dados_edital = json.loads(json_limpo)
                print("\n--- JSON recebido e processado ---")
                print(json.dumps(dados_edital, indent=2, ensure_ascii=False)) # Imprime o JSON de forma legível

                documento = Document(caminho_modelo_docx)
                substituir_formatar(documento, dados_edital)

                documento.save(caminho_arquivo_final)

            except json.JSONDecodeError:
                print("\n--- ERRO DE DECODIFICAÇÃO ---")
                print("A resposta da IA, mesmo após a limpeza, não é um JSON válido. Resposta limpa:")
                print(json_limpo)
        else:
            print("\n--- ERRO ---")
            print("Não foi possível extrair um JSON da resposta da IA. Resposta crua:")
            print(resposta_bruta_ia)

if __name__ == "__main__":
    main()
